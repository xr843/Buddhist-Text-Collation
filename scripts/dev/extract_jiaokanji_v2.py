#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取《顺正理论述文记》第9卷校勘记并生成docx文档
严格按照CBETA原文格式提取，不遗漏、不增删
"""

from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def extract_jiaokanji(file_path):
    """从文件中提取校勘记，严格按照原文格式"""
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    jiaokanji_list = []

    # 用于追踪当前位置的上下文
    current_context = ""
    last_quote = ""

    for i, line in enumerate(lines):
        line_stripped = line.strip()

        # 更新上下文：查找引文（经文或论文）
        # 格式通常是：文字(至)文字。或 文字。
        if line_stripped and not line_stripped.startswith('[') and not line_stripped.startswith('#'):
            if '(至)' in line_stripped or ('。' in line_stripped and len(line_stripped) > 5):
                # 提取引文的前几个字作为位置标识
                clean = re.sub(r'\([^)]*\)', '', line_stripped)
                if len(clean) >= 4 and '述曰' not in clean:
                    last_quote = f"「{clean[:8]}」" if len(clean) > 8 else f"「{clean}」"

        # 匹配校勘记格式: [数字] 或 [数字A/B] 内容
        # 排除 [A数字] 格式（CBETA编者校订）
        match = re.match(r'^\s+\[(\d+[A-B]?)\]\s+(.+)$', line)
        if match:
            num = match.group(1)
            content = match.group(2).strip()

            # 向上查找最近的引文作为位置参考
            context = last_quote
            if not context:
                for j in range(i-1, max(0, i-30), -1):
                    prev = lines[j].strip()
                    if prev and not prev.startswith('[') and not prev.startswith('#'):
                        if '(至)' in prev:
                            clean = re.sub(r'\([^)]*\)', '', prev)
                            context = f"「{clean[:8]}」" if len(clean) > 8 else f"「{clean}」"
                            break
                        elif '述曰' not in prev and len(prev) > 5:
                            context = f"「{prev[:8]}」" if len(prev) > 8 else f"「{prev}」"
                            break

            jiaokanji_list.append({
                'num': num,
                'content': content,
                'context': context
            })

    return jiaokanji_list

def create_docx(jiaokanji_list, output_path):
    """创建docx文档"""
    doc = Document()

    # 设置页面为A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 添加标题
    title = doc.add_heading('《顺正理论述文记》第9卷校勘记', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 添加说明
    intro = doc.add_paragraph()
    intro.add_run(f'共计 {len(jiaokanji_list)} 条校勘记').font.size = Pt(10)
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)

    # 设置表头
    header_cells = table.rows[0].cells
    headers = ['序号', '校勘编号', '校勘内容', '对应《正理论》位置']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加数据行
    for idx, item in enumerate(jiaokanji_list, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = f"[{item['num']}]"
        row_cells[2].text = item['content']
        row_cells[3].text = item['context']

        # 设置单元格格式
        for j, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:
                if j in [0, 1]:  # 序号和编号居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:  # 内容和位置左对齐
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置列宽
    widths = [Cm(1.2), Cm(2), Cm(7), Cm(5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # 保存文档
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")
    return len(jiaokanji_list)

if __name__ == '__main__':
    input_file = '/home/lqsxi/projects/AI-Powered Platform for Buddhist Text Punctuation and Collation Research/CBETA_述文记_卷9.txt'
    output_file = '/home/lqsxi/projects/AI-Powered Platform for Buddhist Text Punctuation and Collation Research/顺正理论述文记第9卷校勘记.docx'

    # 提取校勘记
    jiaokanji = extract_jiaokanji(input_file)

    # 打印前10条检查
    print("前10条校勘记：")
    for i, item in enumerate(jiaokanji[:10], 1):
        print(f"{i}. [{item['num']}] {item['content']} - {item['context']}")

    print(f"\n...")
    print(f"总计: {len(jiaokanji)} 条")

    # 生成文档
    count = create_docx(jiaokanji, output_file)
