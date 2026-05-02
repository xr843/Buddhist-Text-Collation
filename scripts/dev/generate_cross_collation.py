#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从《述文记》卷9校勘记与原文分析，提取可印证《顺正理论》卷12-13校勘的信息
生成校勘记总表docx文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def analyze_cross_collation():
    """
    分析《述文记》卷9中可以印证《顺正理论》卷12-13校勘的内容
    返回校勘条目列表
    """

    # 从《述文记》卷9分析得出的可印证《正理论》校勘的条目
    # 基于述文记原文对正理论的引述和解释

    collation_items = [
        # 卷十二相关
        {
            "seq": 1,
            "volume": "卷12",
            "location": "辯差別品第二之四",
            "zhengliluん_text": "辯【大】，辨【宋】【元】【宮】",
            "shuwenji_evidence": "述文记引「辯差別品」，用「辯」字",
            "analysis": "述文记沿用大正藏底本「辯」字，可印证底本用字",
            "judgment": "大正藏「辯」与四本「辨」为异体字，义同",
            "category": "异体字"
        },
        {
            "seq": 2,
            "volume": "卷12",
            "location": "此已總標",
            "zhengliluん_text": "標【大】，摽【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「此即標列也。一句總標」",
            "analysis": "述文记用「標」字与大正藏一致",
            "judgment": "「標」「摽」为异体字，述文记印证大正藏用字",
            "category": "异体字"
        },
        {
            "seq": 3,
            "volume": "卷12",
            "location": "無色法中已辯心心所",
            "zhengliluん_text": "無",
            "shuwenji_evidence": "述文记云「此下明非色非心法」释「無色法中」",
            "analysis": "述文记解释与正理论原文吻合",
            "judgment": "无异文，印证文本正确",
            "category": "文本印证"
        },
        {
            "seq": 4,
            "volume": "卷12",
            "location": "得謂獲成就",
            "zhengliluん_text": "無",
            "shuwenji_evidence": "述文记详释「一得義中，未得、初得，說名為獲；已得、重得，說名成就」",
            "analysis": "述文记解释与俱舍论比较，印证正理论文本",
            "judgment": "无异文，印证文本正确",
            "category": "文本印证"
        },
        {
            "seq": 5,
            "volume": "卷12",
            "location": "非他相續及非相續",
            "zhengliluん_text": "無",
            "shuwenji_evidence": "述文记云「非他相續，無有成就他身法故；非非相續，無有成就非情法故」",
            "analysis": "述文记逐句释义，与正理论原文一致",
            "judgment": "无异文，印证文本正确",
            "category": "文本印证"
        },
        {
            "seq": 6,
            "volume": "卷12",
            "location": "有太過失",
            "zhengliluん_text": "太【大】，大【元】【明】",
            "shuwenji_evidence": "述文记云「成太過失」",
            "analysis": "述文记用「太」与大正藏一致",
            "judgment": "「太」「大」形近，述文记印证大正藏用「太」",
            "category": "形近致误"
        },
        {
            "seq": 7,
            "volume": "卷12",
            "location": "非他相續及非相續",
            "zhengliluん_text": "二界【大】，三界【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引用此句",
            "analysis": "四本作「三界」，可能是传抄时「二」「三」形近致误",
            "judgment": "需结合义理判断，「二界」或「三界」",
            "category": "形近致误"
        },
        {
            "seq": 8,
            "volume": "卷12",
            "location": "彼時具證眼等結斷",
            "zhengliluん_text": "彼【大】，後【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记相关段落未见此字",
            "analysis": "「彼」「後」字形相近，四本异文",
            "judgment": "形近致误，需据义理判断",
            "category": "形近致误"
        },
        {
            "seq": 9,
            "volume": "卷12",
            "location": "生性等例",
            "zhengliluん_text": "生【大】，牛【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「由此已遮生性等例」",
            "analysis": "述文记用「生」，与大正藏一致",
            "judgment": "「生」「牛」形近致误，述文记印证大正藏「生」为正",
            "category": "形近致误"
        },
        {
            "seq": 10,
            "volume": "卷12",
            "location": "蒃豆等生",
            "zhengliluん_text": "菉【大】，綠【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "「菉」「綠」为异体字",
            "judgment": "异体字，义同",
            "category": "异体字"
        },
        {
            "seq": 11,
            "volume": "卷12",
            "location": "色廛有情",
            "zhengliluん_text": "廛【大】，纏【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记卷13亦见「廛」字",
            "analysis": "「廛」「纏」音近，四本作「纏」",
            "judgment": "同音假借或传抄讹误",
            "category": "同音假借"
        },
        # 卷十三相关
        {
            "seq": 12,
            "volume": "卷13",
            "location": "五蘊中已廣分別",
            "zhengliluん_text": "五【大】，立【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记释「如五蘊中」",
            "analysis": "「五」与「立」字形差异大，四本异文待考",
            "judgment": "可能为传抄讹误，「五」义胜",
            "category": "实质异文"
        },
        {
            "seq": 13,
            "volume": "卷13",
            "location": "故偏厭逆",
            "zhengliluん_text": "偏【大】，遍【宮】",
            "shuwenji_evidence": "述文记云「由方便中漏厭此二」",
            "analysis": "宫本作「遍」，与大正藏「偏」不同",
            "judgment": "「偏」「遍」音近，大正藏「偏」义胜",
            "category": "同音假借"
        },
        {
            "seq": 14,
            "volume": "卷13",
            "location": "正理皆不相違",
            "zhengliluん_text": "正【大】，王【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "「正」「王」形近，四本皆作「王」",
            "judgment": "形近致误，「正理」为专有名词，当作「正」",
            "category": "形近致误"
        },
        {
            "seq": 15,
            "volume": "卷13",
            "location": "謂言不可知便非相者",
            "zhengliluん_text": "謂【大】，言【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本作「言」，与大正藏「謂」异",
            "judgment": "「謂」「言」义近，皆可通",
            "category": "实质异文"
        },
        {
            "seq": 16,
            "volume": "卷13",
            "location": "辯差別品",
            "zhengliluん_text": "辯【大】，辨【宋】【元】【宮】",
            "shuwenji_evidence": "述文记用「辯」字",
            "analysis": "述文记与大正藏一致用「辯」",
            "judgment": "「辯」「辨」异体字，义同",
            "category": "异体字"
        },
        {
            "seq": 17,
            "volume": "卷13",
            "location": "故【CB】【麗-CB】，果【大】，起【宋】【元】【明】【宮】",
            "zhengliluん_text": "故【CB】，果【大】，起【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「無異因故」",
            "analysis": "CBETA据高丽藏改「果」为「故」",
            "judgment": "「故」义胜，述文记用「故」印证CBETA校订",
            "category": "CBETA校订"
        },
        {
            "seq": 18,
            "volume": "卷13",
            "location": "羯剌藍",
            "zhengliluん_text": "剌【CB】【麗-CB】【宋】【元】【明】【宮】，刺【大】",
            "shuwenji_evidence": "述文记云「羯剌藍」",
            "analysis": "大正藏作「刺」，诸本作「剌」",
            "judgment": "「剌」为正字，大正藏「刺」为形近讹误",
            "category": "形近致误"
        },
        {
            "seq": 19,
            "volume": "卷13",
            "location": "鳩磨羅時",
            "zhengliluん_text": "磨【大】，摩【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「鳩摩羅時婆」",
            "analysis": "述文记用「摩」与四本一致",
            "judgment": "「摩」为正字，述文记印证四本",
            "category": "异体字"
        },
        {
            "seq": 20,
            "volume": "卷13",
            "location": "彼經舉後",
            "zhengliluん_text": "彼【大】，婆【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "「彼」「婆」形近，四本异文",
            "judgment": "「彼」义胜，四本「婆」为形近讹误",
            "category": "形近致误"
        },
        {
            "seq": 21,
            "volume": "卷13",
            "location": "輕大",
            "zhengliluん_text": "經【CB】【麗-CB】【宋】【元】【明】【宮】，輕【大】",
            "shuwenji_evidence": "述文记引「契經」",
            "analysis": "大正藏作「輕」，诸本作「經」",
            "judgment": "「經」为正字，CBETA校订正确",
            "category": "CBETA校订"
        },
        {
            "seq": 22,
            "volume": "卷13",
            "location": "必為宋元明宮",
            "zhengliluん_text": "必【大】，為【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本皆作「為」，与大正藏「必」异",
            "judgment": "需据上下文判断，「必」「為」义别",
            "category": "实质异文"
        },
        {
            "seq": 23,
            "volume": "卷13",
            "location": "功能差別",
            "zhengliluん_text": "功【大】，切【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记多次引「功能」",
            "analysis": "四本作「切」，与大正藏「功」异",
            "judgment": "「功能」为专有术语，大正藏「功」为正",
            "category": "形近致误"
        },
        {
            "seq": 24,
            "volume": "卷13",
            "location": "能餘宋元明宮",
            "zhengliluん_text": "能【大】，餘【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本皆作「餘」，与大正藏「能」异",
            "judgment": "需据上下文判断",
            "category": "实质异文"
        },
        {
            "seq": 25,
            "volume": "卷13",
            "location": "便更宋元明宮",
            "zhengliluん_text": "便【大】，更【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本皆作「更」，与大正藏「便」异",
            "judgment": "「便」「更」义近，皆可通",
            "category": "实质异文"
        },
        {
            "seq": 26,
            "volume": "卷13",
            "location": "摽宋元明宮",
            "zhengliluん_text": "標【大】，摽【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记用「標」",
            "analysis": "与卷12同，述文记用「標」",
            "judgment": "「標」「摽」异体字，义同",
            "category": "异体字"
        },
        {
            "seq": 27,
            "volume": "卷13",
            "location": "亦應應亦",
            "zhengliluん_text": "亦應【大】，應亦【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本作「應亦」，与大正藏「亦應」语序不同",
            "judgment": "语序倒置，义同",
            "category": "语序异文"
        },
        {
            "seq": 28,
            "volume": "卷13",
            "location": "羯賴藍",
            "zhengliluん_text": "賴【大】，剌【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「羯剌藍」",
            "analysis": "述文记用「剌」与四本一致",
            "judgment": "「剌」为正字，述文记印证四本",
            "category": "形近致误"
        },
        {
            "seq": 29,
            "volume": "卷13",
            "location": "過遇",
            "zhengliluん_text": "過【大】，遇【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记云「由過未得」",
            "analysis": "述文记用「過」与大正藏一致",
            "judgment": "「過」「遇」形近，大正藏「過」义胜",
            "category": "形近致误"
        },
        {
            "seq": 30,
            "volume": "卷13",
            "location": "辦辨辯",
            "zhengliluん_text": "辦【大】，辨【宋】【元】【宮】，辯【明】",
            "shuwenji_evidence": "述文记用「辨」",
            "analysis": "「辦」「辨」「辯」三字混用",
            "judgment": "异体字，义同",
            "category": "异体字"
        },
        {
            "seq": 31,
            "volume": "卷13",
            "location": "果異",
            "zhengliluん_text": "果【大】，異【宋】【元】【明】【宮】",
            "shuwenji_evidence": "述文记未直接引此句",
            "analysis": "四本皆作「異」，与大正藏「果」异",
            "judgment": "需据上下文判断，「異」「果」义别",
            "category": "实质异文"
        },
    ]

    return collation_items

def create_statistics(items):
    """创建统计数据"""
    stats = {
        "total": len(items),
        "vol12": len([i for i in items if i["volume"] == "卷12"]),
        "vol13": len([i for i in items if i["volume"] == "卷13"]),
        "categories": {}
    }

    for item in items:
        cat = item["category"]
        if cat not in stats["categories"]:
            stats["categories"][cat] = 0
        stats["categories"][cat] += 1

    return stats

def create_docx(items, stats, output_path):
    """创建docx文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    title = doc.add_heading('《顺正理论》卷12-13校勘记总表', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.add_run('——基于《述文记》卷9的印证分析').font.size = Pt(12)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 说明
    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run('说明：').bold = True
    intro.add_run('本表通过分析《顺正理论述文记》第9卷原文及其校勘记，提取可印证《顺正理论》卷12-13校勘的信息。《述文记》作为注疏文献，其引文和解释可从另一角度验证原典文本的正确性。')

    doc.add_paragraph()

    # 统计表
    doc.add_heading('一、统计概览', level=2)

    stat_table = doc.add_table(rows=1, cols=2)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(stat_table)

    # 表头
    header = stat_table.rows[0].cells
    header[0].text = "统计项目"
    header[1].text = "数量"
    for cell in header:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True

    # 统计数据
    stat_data = [
        ("校勘条目总数", str(stats["total"])),
        ("卷十二相关条目", str(stats["vol12"])),
        ("卷十三相关条目", str(stats["vol13"])),
    ]
    for cat, count in stats["categories"].items():
        stat_data.append((f"  - {cat}", str(count)))

    for label, value in stat_data:
        row = stat_table.add_row().cells
        row[0].text = label
        row[1].text = value
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 校勘记总表
    doc.add_heading('二、校勘记总表', level=2)

    main_table = doc.add_table(rows=1, cols=6)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(main_table)

    # 表头
    headers = ['序号', '卷次', '《正理论》校勘', '《述文记》印证', '判取意见', '类型']
    header_row = main_table.rows[0].cells
    for i, h in enumerate(headers):
        header_row[i].text = h
        for p in header_row[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for item in items:
        row = main_table.add_row().cells
        row[0].text = str(item["seq"])
        row[1].text = item["volume"]
        row[2].text = item["zhengliluん_text"]
        row[3].text = item["shuwenji_evidence"]
        row[4].text = item["judgment"]
        row[5].text = item["category"]

        for j, cell in enumerate(row):
            for p in cell.paragraphs:
                if j in [0, 1, 5]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)

    # 设置列宽
    widths = [Cm(1), Cm(1.5), Cm(4), Cm(4), Cm(4), Cm(2)]
    for row in main_table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    doc.add_paragraph()

    # 结论
    doc.add_heading('三、分析结论', level=2)

    conclusion_items = [
        "1. 异体字问题：「辯/辨/辦」「標/摽」「菉/綠」等为常见异体字，不影响文义理解。",
        "2. 形近致误：「太/大」「生/牛」「彼/後」「正/王」「刺/剌」等形近字在传抄中容易混淆。",
        "3. 述文记印证价值：述文记作为注疏，其引文可印证原典用字，如「生性」「太過」等。",
        "4. CBETA校订：部分CBETA校订（如「故/果」「經/輕」）可从述文记得到印证。",
        "5. 实质异文：部分异文涉及义理差别，需结合上下文审慎判断。",
    ]

    for item in conclusion_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.first_line_indent = Cm(0.7)

    # 保存
    doc.save(output_path)
    print(f"文档已保存至: {output_path}")

if __name__ == '__main__':
    # 分析校勘信息
    items = analyze_cross_collation()

    # 统计
    stats = create_statistics(items)

    print(f"总条目数: {stats['total']}")
    print(f"卷12条目: {stats['vol12']}")
    print(f"卷13条目: {stats['vol13']}")
    print("分类统计:")
    for cat, count in stats['categories'].items():
        print(f"  {cat}: {count}")

    # 生成文档
    output_path = '/home/lqsxi/projects/AI-Powered Platform for Buddhist Text Punctuation and Collation Research/顺正理论卷12-13校勘记总表_述文记印证.docx'
    create_docx(items, stats, output_path)
