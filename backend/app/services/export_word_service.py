"""
Word格式导出服务
生成符合学术规范的校勘报告
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Any


class WordExportService:
    """Word导出服务"""

    @staticmethod
    def export_collation_report(
        collation_data: Dict[str, Any],
        report_type: str = "multi_version"  # two_version 或 multi_version
    ) -> BytesIO:
        """
        导出校勘报告为Word格式

        Args:
            collation_data: 校勘数据
                {
                    'base_name': str,
                    'base_text': str,
                    'collations': List[{'name': str, 'text': str}],
                    'differences': List[{
                        'position': int,
                        'context': str,
                        'base_char': str,
                        'variants': List[{'name': str, 'char': str, 'category': str}],
                        'category': str,
                        'note': str
                    }],
                    'statistics': {...}
                }
            report_type: 报告类型

        Returns:
            BytesIO: Word文档字节流
        """
        doc = Document()

        # 设置中文字体
        WordExportService._set_chinese_font(doc)

        # 1. 标题
        title = doc.add_heading('校勘记', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = '宋体'
        title_run.font.size = Pt(22)
        title_run.font.bold = True

        # 添加空行
        doc.add_paragraph()

        # 2. 基本信息
        WordExportService._add_basic_info(doc, collation_data)

        # 3. 统计信息
        if 'statistics' in collation_data:
            WordExportService._add_statistics(doc, collation_data['statistics'])

        # 4. 异文汇校表
        if report_type == "multi_version":
            WordExportService._add_multi_version_table(doc, collation_data)
        else:
            WordExportService._add_two_version_table(doc, collation_data)

        # 5. 页脚（生成信息）
        WordExportService._add_footer(doc)

        # 保存到字节流
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _set_chinese_font(doc: Document):
        """设置中文字体支持"""
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    @staticmethod
    def _add_basic_info(doc: Document, data: Dict):
        """添加基本信息"""
        info_section = doc.add_paragraph()
        info_section.add_run('一、基本信息\n').bold = True

        # 底本信息
        p1 = doc.add_paragraph()
        p1.add_run('底本：').bold = True
        p1.add_run(data.get('base_name', '未命名'))

        # 校本信息
        collations = data.get('collations', [])
        if collations:
            p2 = doc.add_paragraph()
            p2.add_run('校本：').bold = True
            collation_names = [c.get('name', '未命名') for c in collations]
            p2.add_run(', '.join(collation_names))

        # 异文总数
        differences = data.get('differences', [])
        p3 = doc.add_paragraph()
        p3.add_run('异文总数：').bold = True
        p3.add_run(str(len(differences)))

        # 生成时间
        p4 = doc.add_paragraph()
        p4.add_run('生成时间：').bold = True
        p4.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))

        doc.add_paragraph()  # 空行

    @staticmethod
    def _add_statistics(doc: Document, stats: Dict):
        """添加统计信息"""
        stats_section = doc.add_paragraph()
        stats_section.add_run('二、统计分析\n').bold = True

        # 按类型统计
        if 'by_category' in stats:
            p = doc.add_paragraph()
            p.add_run('异文类型分布：\n').bold = True
            for category, count in stats['by_category'].items():
                category_names = {
                    'error': '讹误',
                    'variant': '异体字',
                    'yanwen': '衍文',
                    'tuowen': '脱文',
                    'daowen': '倒文'
                }
                cat_name = category_names.get(category, category)
                doc.add_paragraph(f'  • {cat_name}：{count}处', style='List Bullet')

        # 相似度（如果有）
        if 'similarity' in stats:
            p = doc.add_paragraph()
            p.add_run('相似度：').bold = True
            p.add_run(f"{stats['similarity']:.1f}%")

        doc.add_paragraph()  # 空行

    @staticmethod
    def _add_multi_version_table(doc: Document, data: Dict):
        """添加多版本对勘表"""
        table_section = doc.add_paragraph()
        table_section.add_run('三、异文汇校表\n').bold = True

        differences = data.get('differences', [])
        if not differences:
            doc.add_paragraph('（无异文）')
            return

        collations = data.get('collations', [])

        # 创建表格：序号、上下文、底本、各校本、类型
        col_count = 4 + len(collations)  # 序号、上下文、底本、N个校本、类型
        table = doc.add_table(rows=1, cols=col_count)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        header_cells = table.rows[0].cells
        headers = ['序号', '上下文', '底本'] + [c.get('name', f'校本{i+1}') for i, c in enumerate(collations)] + ['类型']

        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 设置背景色
            WordExportService._set_cell_background(cell, 'D3D3D3')

        # 数据行
        for idx, diff in enumerate(differences[:500], 1):  # 限制最多500条，避免文档过大
            row = table.add_row().cells

            # 序号
            row[0].text = str(idx)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 上下文
            context = diff.get('context', '...')[:20]  # 限制长度
            row[1].text = context

            # 底本
            base_char = diff.get('base_char', '')
            WordExportService._add_colored_text(row[2], base_char, '1890ff')
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 各校本
            variants_by_name = {v['name']: v for v in diff.get('variants', [])}
            for i, collation in enumerate(collations):
                cell_idx = 3 + i
                variant = variants_by_name.get(collation['name'])
                if variant:
                    char = variant.get('char', '')
                    category = variant.get('category', '')

                    # 根据类型设置颜色
                    color = WordExportService._get_category_color(category)
                    WordExportService._add_colored_text(row[cell_idx], char, color)
                else:
                    row[cell_idx].text = '同底本'
                row[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 类型
            category = diff.get('category', '')
            category_names = {
                'error': '讹误',
                'variant': '异体字',
                'yanwen': '衍文',
                'tuowen': '脱文',
                'daowen': '倒文'
            }
            row[-1].text = category_names.get(category, category)
            row[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 如果超过500条，添加说明
        if len(differences) > 500:
            doc.add_paragraph()
            note = doc.add_paragraph()
            note.add_run(f'注：异文总数{len(differences)}条，此处仅展示前500条。完整数据请参考CSV导出。').italic = True
            note.runs[0].font.size = Pt(10)
            note.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    @staticmethod
    def _add_two_version_table(doc: Document, data: Dict):
        """添加两版本对勘表（简化版）"""
        table_section = doc.add_paragraph()
        table_section.add_run('三、异文汇校表\n').bold = True

        differences = data.get('differences', [])
        if not differences:
            doc.add_paragraph('（无异文）')
            return

        # 创建表格：序号、位置、上下文、底本、校本、类型、说明
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'

        # 表头
        headers = ['序号', '位置', '上下文', '底本', '校本', '类型', '说明']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            WordExportService._set_cell_background(cell, 'D3D3D3')

        # 数据行
        for idx, diff in enumerate(differences[:500], 1):
            row = table.add_row().cells

            row[0].text = str(idx)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row[1].text = str(diff.get('position', ''))
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row[2].text = diff.get('context', '')[:20]

            row[3].text = diff.get('base_char', '')
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 校本（取第一个variant）
            variants = diff.get('variants', [])
            if variants:
                row[4].text = variants[0].get('char', '')
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            category = diff.get('category', '')
            category_names = {
                'error': '讹误',
                'variant': '异体字',
                'yanwen': '衍文',
                'tuowen': '脱文',
                'daowen': '倒文'
            }
            row[5].text = category_names.get(category, category)
            row[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            row[6].text = diff.get('note', '')

    @staticmethod
    def _add_footer(doc: Document):
        """添加页脚"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f'本报告由"佛典标点与校勘研究平台"生成于{datetime.now().strftime("%Y年%m月%d日")}'
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = Pt(9)
        footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    @staticmethod
    def _set_cell_background(cell, color_hex: str):
        """设置单元格背景色"""
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._element.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _add_colored_text(cell, text: str, color_hex: str):
        """添加彩色文本"""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
        run.font.bold = True

    @staticmethod
    def _get_category_color(category: str) -> str:
        """获取类别对应的颜色（十六进制）"""
        color_map = {
            'error': 'ff0000',      # 红色 - 讹误
            'variant': '00aa00',    # 绿色 - 异体字
            'yanwen': 'ff8800',     # 橙色 - 衍文
            'tuowen': 'ff8800',     # 橙色 - 脱文
            'daowen': 'aa00ff',     # 紫色 - 倒文
        }
        return color_map.get(category, '000000')  # 默认黑色


    @staticmethod
    def export_phylogeny_report(phylogeny_data: Dict[str, Any]) -> BytesIO:
        """
        导出版本谱系分析报告为Word格式

        Args:
            phylogeny_data: 谱系分析数据
                {
                    'similarity_matrix': {...},
                    'shared_errors': {...},
                    'tree': {...},
                    'conclusions': [...]
                }
        """
        doc = Document()
        WordExportService._set_chinese_font(doc)

        # 标题
        title = doc.add_heading('版本谱系分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 基本信息
        similarity_matrix = phylogeny_data.get('similarity_matrix', {})
        names = similarity_matrix.get('names', [])

        p1 = doc.add_paragraph()
        p1.add_run('底本：').bold = True
        p1.add_run(names[0] if names else '未指定')

        p2 = doc.add_paragraph()
        p2.add_run('校本数量：').bold = True
        p2.add_run(str(len(names) - 1 if len(names) > 0 else 0))

        p3 = doc.add_paragraph()
        p3.add_run('生成时间：').bold = True
        p3.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))

        doc.add_paragraph()

        # 分析结论
        conclusions = phylogeny_data.get('conclusions', [])
        if conclusions:
            doc.add_heading('分析结论', level=2)
            for i, conclusion in enumerate(conclusions, 1):
                doc.add_paragraph(f'{i}. {conclusion}', style='List Number')
            doc.add_paragraph()

        # 相似度矩阵
        doc.add_heading('相似度矩阵', level=2)
        matrix = similarity_matrix.get('matrix', [])
        if matrix and names:
            # 创建表格
            table = doc.add_table(rows=len(names)+1, cols=len(names)+1)
            table.style = 'Light Grid Accent 1'

            # 表头
            table.rows[0].cells[0].text = ''
            for i, name in enumerate(names):
                table.rows[0].cells[i+1].text = name[:10]  # 限制长度
                table.rows[0].cells[i+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 数据
            for i, name in enumerate(names):
                table.rows[i+1].cells[0].text = name[:10]
                for j, sim in enumerate(matrix[i]):
                    cell = table.rows[i+1].cells[j+1]
                    if i == j:
                        cell.text = '-'
                    else:
                        cell.text = f'{sim*100:.1f}%'
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 页脚
        WordExportService._add_footer(doc)

        # 保存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def export_collation_notes_report(
        notes: List[Dict[str, Any]],
        title: str = "校勘记",
        base_name: str = "底本",
        collation_names: List[str] = None,
        include_statistics: bool = True,
        include_base_info: bool = True
    ) -> BytesIO:
        """
        导出校勘记报告为Word格式

        Args:
            notes: 校勘记数据列表
                [
                    {
                        'index': 1,
                        'position': 23,
                        'position_display': '第24字',
                        'original_char': '住',
                        'replacement_char': '往',
                        'action': '改',
                        'source_versions': ['丽藏'],
                        'explanation': '按：多数版本作「往」。',
                        'category': '讹误',
                        'uncertain': False,
                        'formatted_text': '【1】第24字「住」，据丽藏改「往」。按：多数版本作「往」。'
                    },
                    ...
                ]
            title: 报告标题
            base_name: 底本名称
            collation_names: 校本名称列表
            include_statistics: 是否包含统计信息
            include_base_info: 是否包含基本信息

        Returns:
            BytesIO: Word文档字节流
        """
        doc = Document()
        collation_names = collation_names or []

        # 设置中文字体
        WordExportService._set_chinese_font(doc)

        # 1. 标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.name = '宋体'
        heading.runs[0].font.size = Pt(22)
        heading.runs[0].font.bold = True

        doc.add_paragraph()

        # 2. 基本信息
        if include_base_info:
            info_section = doc.add_paragraph()
            info_section.add_run('一、基本信息\n').bold = True

            p1 = doc.add_paragraph()
            p1.add_run('底本：').bold = True
            p1.add_run(base_name)

            if collation_names:
                p2 = doc.add_paragraph()
                p2.add_run('校本：').bold = True
                p2.add_run('、'.join(collation_names))

            p3 = doc.add_paragraph()
            p3.add_run('校勘记条数：').bold = True
            p3.add_run(str(len(notes)))

            p4 = doc.add_paragraph()
            p4.add_run('生成时间：').bold = True
            p4.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))

            doc.add_paragraph()

        # 3. 统计信息
        if include_statistics and notes:
            stats_section = doc.add_paragraph()
            stats_section.add_run('二、统计分析\n').bold = True

            # 按类型统计
            by_category: Dict[str, int] = {}
            for note in notes:
                cat = note.get('category', '其他')
                by_category[cat] = by_category.get(cat, 0) + 1

            p = doc.add_paragraph()
            p.add_run('异文类型分布：\n').bold = True
            for category, count in by_category.items():
                doc.add_paragraph(f'  • {category}：{count}处', style='List Bullet')

            # 存疑统计
            uncertain_count = sum(1 for n in notes if n.get('uncertain', False))
            if uncertain_count > 0:
                p_uncertain = doc.add_paragraph()
                p_uncertain.add_run('存疑条目：').bold = True
                p_uncertain.add_run(f'{uncertain_count}处')

            doc.add_paragraph()

        # 4. 校勘记正文
        notes_section = doc.add_paragraph()
        section_title = '三、校勘记\n' if include_base_info else '校勘记\n'
        notes_section.add_run(section_title).bold = True

        for note in notes:
            p = doc.add_paragraph()
            formatted_text = note.get('formatted_text', '')

            # 存疑项标橙色
            if note.get('uncertain', False):
                run = p.add_run(formatted_text)
                run.font.color.rgb = RGBColor(255, 128, 0)  # 橙色
            else:
                p.add_run(formatted_text)

        # 5. 页脚
        WordExportService._add_footer(doc)

        # 保存到字节流
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
