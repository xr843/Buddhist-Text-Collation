"""
校勘记生成与导出路由

提供校勘记生成、获取、编辑和导出的API端点。
"""
from fastapi import APIRouter, HTTPException, status
from fastapi.responses import StreamingResponse
from typing import Optional, List, Dict, Any
from pydantic import BaseModel, Field
from io import BytesIO
from datetime import datetime

from app.services.project_storage import project_storage, ProjectType
from app.services.collation_note_generator import collation_note_generator, CollationNote

router = APIRouter()


# ============ 请求/响应模型 ============

class GenerateNotesRequest(BaseModel):
    """生成校勘记请求"""
    regenerate: bool = False  # 是否重新生成（覆盖已有）


class GenerateNotesResponse(BaseModel):
    """生成校勘记响应"""
    success: bool
    project_id: str
    notes: List[Dict[str, Any]]
    total: int
    message: str


class UpdateNoteRequest(BaseModel):
    """更新校勘记请求"""
    formatted_text: Optional[str] = None
    explanation: Optional[str] = None
    action: Optional[str] = None
    uncertain: Optional[bool] = None


class ExportNotesRequest(BaseModel):
    """导出校勘记请求"""
    format: str = Field(default="word", description="导出格式：word/txt/markdown")
    include_statistics: bool = True
    include_base_info: bool = True


# ============ API端点 ============

@router.post("/projects/{project_id}/generate-collation-notes")
async def generate_collation_notes(
    project_id: str,
    request: GenerateNotesRequest = GenerateNotesRequest()
) -> GenerateNotesResponse:
    """
    生成校勘记

    基于项目的判取结果，自动生成符合《中华大藏经》体例的校勘记。

    Args:
        project_id: 项目ID
        request: 生成请求参数

    Returns:
        生成的校勘记列表
    """
    try:
        project = project_storage.get_project(ProjectType.MULTI_COLLATION, project_id)
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"项目不存在: {project_id}"
            )

        # 获取判取结果
        decisions = project.get("data", {}).get("decisions", {})
        if not decisions:
            return GenerateNotesResponse(
                success=True,
                project_id=project_id,
                notes=[],
                total=0,
                message="无判取结果，无法生成校勘记"
            )

        # 获取底本信息
        base_data = project["data"].get("base", {})
        base_text = base_data.get("text", "")
        base_name = base_data.get("name", "底本")

        # 获取校本名称列表
        collations = project["data"].get("collations", [])
        collation_names = [c.get("name", f"校本{i+1}") for i, c in enumerate(collations)]

        # 获取异文表（如果有）- 注意数据结构是 {headers, rows, total}
        variant_table_data = project["data"].get("variant_table", {})
        if isinstance(variant_table_data, dict):
            variant_table = variant_table_data.get("rows", [])
        else:
            variant_table = variant_table_data if isinstance(variant_table_data, list) else []

        # 生成校勘记
        notes = collation_note_generator.generate_notes(
            decisions=decisions,
            base_text=base_text,
            base_name=base_name,
            collation_names=collation_names,
            variant_table=variant_table
        )

        # 转换为字典格式
        notes_data = [note.model_dump() for note in notes]

        # 保存到项目
        project["data"]["collation_notes"] = notes_data
        project["data"]["collation_notes_generated_at"] = datetime.now().isoformat()

        project_storage.update_project(
            ProjectType.MULTI_COLLATION,
            project_id,
            data=project["data"],
            merge_data=False
        )

        return GenerateNotesResponse(
            success=True,
            project_id=project_id,
            notes=notes_data,
            total=len(notes_data),
            message=f"成功生成 {len(notes_data)} 条校勘记"
        )

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[生成校勘记] 错误详情: {type(e).__name__}: {str(e)}")
        print(f"[生成校勘记] 完整堆栈:\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"生成校勘记失败: {str(e)}"
        )


@router.get("/projects/{project_id}/collation-notes")
async def get_collation_notes(project_id: str):
    """
    获取已生成的校勘记

    Args:
        project_id: 项目ID

    Returns:
        校勘记列表和相关信息
    """
    try:
        project = project_storage.get_project(ProjectType.MULTI_COLLATION, project_id)
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"项目不存在: {project_id}"
            )

        notes = project.get("data", {}).get("collation_notes", [])
        generated_at = project.get("data", {}).get("collation_notes_generated_at")

        # 统计信息
        total = len(notes)
        uncertain_count = sum(1 for n in notes if n.get("uncertain", False))
        decided_count = total - uncertain_count

        # 按类型统计
        by_category = {}
        for note in notes:
            cat = note.get("category", "其他")
            by_category[cat] = by_category.get(cat, 0) + 1

        # 按动作统计
        by_action = {}
        for note in notes:
            action = note.get("action", "其他")
            by_action[action] = by_action.get(action, 0) + 1

        return {
            "success": True,
            "project_id": project_id,
            "notes": notes,
            "total": total,
            "generated_at": generated_at,
            "statistics": {
                "total": total,
                "decided": decided_count,
                "uncertain": uncertain_count,
                "by_category": by_category,
                "by_action": by_action
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取校勘记失败: {str(e)}"
        )


@router.put("/projects/{project_id}/collation-notes/{note_index}")
async def update_collation_note(
    project_id: str,
    note_index: int,
    request: UpdateNoteRequest
):
    """
    编辑单条校勘记

    允许用户修改校勘记的文本、说明等内容。

    Args:
        project_id: 项目ID
        note_index: 校勘记序号（1-based）
        request: 更新内容

    Returns:
        更新后的校勘记
    """
    try:
        project = project_storage.get_project(ProjectType.MULTI_COLLATION, project_id)
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"项目不存在: {project_id}"
            )

        notes = project.get("data", {}).get("collation_notes", [])

        # 找到对应的校勘记（note_index是1-based，对应note.index字段）
        note_idx = None
        for i, note in enumerate(notes):
            if note.get("index") == note_index:
                note_idx = i
                break

        if note_idx is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"校勘记不存在: 序号 {note_index}"
            )

        # 更新字段
        note = notes[note_idx]
        if request.formatted_text is not None:
            note["formatted_text"] = request.formatted_text
        if request.explanation is not None:
            note["explanation"] = request.explanation
        if request.action is not None:
            note["action"] = request.action
        if request.uncertain is not None:
            note["uncertain"] = request.uncertain

        note["updated_at"] = datetime.now().isoformat()

        # 保存
        project["data"]["collation_notes"] = notes
        project_storage.update_project(
            ProjectType.MULTI_COLLATION,
            project_id,
            data=project["data"],
            merge_data=False
        )

        return {
            "success": True,
            "message": f"已更新校勘记 {note_index}",
            "note": note
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"更新校勘记失败: {str(e)}"
        )


@router.delete("/projects/{project_id}/collation-notes/{note_index}")
async def delete_collation_note(
    project_id: str,
    note_index: int,
    delete_decision: bool = False
):
    """
    删除单条校勘记

    删除后会自动重新编号。

    Args:
        project_id: 项目ID
        note_index: 校勘记序号（1-based）
        delete_decision: 是否同时删除对应的判取记录

    Returns:
        删除结果
    """
    try:
        project = project_storage.get_project(ProjectType.MULTI_COLLATION, project_id)
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"项目不存在: {project_id}"
            )

        notes = project.get("data", {}).get("collation_notes", [])

        # 找到对应的校勘记
        note_idx = None
        for i, note in enumerate(notes):
            if note.get("index") == note_index:
                note_idx = i
                break

        if note_idx is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"校勘记不存在: 序号 {note_index}"
            )

        # 删除校勘记
        deleted_note = notes.pop(note_idx)

        # 如果需要同时删除判取记录
        deleted_decision_position = None
        if delete_decision:
            position = deleted_note.get("position")
            if position is not None:
                decisions = project.get("data", {}).get("decisions", {})
                position_str = str(position)
                if position_str in decisions:
                    del decisions[position_str]
                    project["data"]["decisions"] = decisions
                    deleted_decision_position = position

        # 重新编号
        for i, note in enumerate(notes):
            note["index"] = i + 1

        # 保存
        project["data"]["collation_notes"] = notes
        project_storage.update_project(
            ProjectType.MULTI_COLLATION,
            project_id,
            data=project["data"],
            merge_data=False
        )

        return {
            "success": True,
            "message": f"已删除校勘记 {note_index}" + (f"，同时删除了位置 {deleted_decision_position} 的判取记录" if deleted_decision_position is not None else ""),
            "deleted_note": deleted_note,
            "deleted_decision_position": deleted_decision_position,
            "remaining_count": len(notes)
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"删除校勘记失败: {str(e)}"
        )


@router.post("/projects/{project_id}/export-collation-notes")
async def export_collation_notes(
    project_id: str,
    request: ExportNotesRequest = ExportNotesRequest()
):
    """
    导出校勘记

    支持导出为Word、TXT、Markdown格式。

    Args:
        project_id: 项目ID
        request: 导出请求参数

    Returns:
        文件下载响应
    """
    try:
        project = project_storage.get_project(ProjectType.MULTI_COLLATION, project_id)
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"项目不存在: {project_id}"
            )

        notes = project.get("data", {}).get("collation_notes", [])
        if not notes:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="无校勘记可导出，请先生成校勘记"
            )

        # 获取项目信息
        title = project.get("title", "校勘记")
        base_name = project.get("data", {}).get("base", {}).get("name", "底本")
        collations = project.get("data", {}).get("collations", [])
        collation_names = [c.get("name", f"校本{i+1}") for i, c in enumerate(collations)]

        export_format = request.format.lower()

        if export_format == "word":
            return await _export_notes_word(
                notes, title, base_name, collation_names,
                request.include_statistics, request.include_base_info
            )
        elif export_format == "markdown" or export_format == "md":
            return await _export_notes_markdown(
                notes, title, base_name, collation_names,
                request.include_statistics, request.include_base_info
            )
        elif export_format == "txt" or export_format == "text":
            return await _export_notes_txt(
                notes, title, base_name, collation_names,
                request.include_statistics, request.include_base_info
            )
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"不支持的导出格式: {export_format}，支持: word/txt/markdown"
            )

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[导出校勘记] 错误详情: {type(e).__name__}: {str(e)}")
        print(f"[导出校勘记] 完整堆栈:\n{traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"导出校勘记失败: {str(e)}"
        )


# ============ 导出辅助函数 ============

async def _export_notes_word(
    notes: List[Dict],
    title: str,
    base_name: str,
    collation_names: List[str],
    include_statistics: bool,
    include_base_info: bool
) -> StreamingResponse:
    """导出Word格式校勘记"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    heading = doc.add_heading('校勘记', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.name = '宋体'
    heading.runs[0].font.size = Pt(22)

    doc.add_paragraph()

    # 基本信息
    if include_base_info:
        info_section = doc.add_paragraph()
        info_section.add_run('一、基本信息\n').bold = True

        p1 = doc.add_paragraph()
        p1.add_run('底本：').bold = True
        p1.add_run(base_name)

        if collation_names:
            p2 = doc.add_paragraph()
            p2.add_run('校本：').bold = True
            p2.add_run('、'.join(collation_names))

        p3 = doc.add_paragraph()
        p3.add_run('校勘记条数：').bold = True
        p3.add_run(str(len(notes)))

        p4 = doc.add_paragraph()
        p4.add_run('生成时间：').bold = True
        p4.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))

        doc.add_paragraph()

    # 统计信息
    if include_statistics:
        stats_section = doc.add_paragraph()
        stats_section.add_run('二、统计分析\n').bold = True

        # 按类型统计
        by_category = {}
        for note in notes:
            cat = note.get("category", "其他")
            by_category[cat] = by_category.get(cat, 0) + 1

        p = doc.add_paragraph()
        p.add_run('异文类型分布：\n').bold = True
        for category, count in by_category.items():
            doc.add_paragraph(f'  • {category}：{count}处', style='List Bullet')

        # 存疑统计
        uncertain_count = sum(1 for n in notes if n.get("uncertain", False))
        if uncertain_count > 0:
            p_uncertain = doc.add_paragraph()
            p_uncertain.add_run('存疑条目：').bold = True
            p_uncertain.add_run(f'{uncertain_count}处')

        doc.add_paragraph()

    # 校勘记正文
    notes_section = doc.add_paragraph()
    notes_section.add_run('三、校勘记\n' if include_base_info else '校勘记\n').bold = True

    for note in notes:
        p = doc.add_paragraph()
        formatted_text = note.get("formatted_text", "")

        # 存疑项标红
        if note.get("uncertain", False):
            run = p.add_run(formatted_text)
            run.font.color.rgb = RGBColor(255, 128, 0)  # 橙色
        else:
            p.add_run(formatted_text)

    # 页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'本报告由"佛典标点与校勘研究平台"生成于{datetime.now().strftime("%Y年%m月%d日")}'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # 保存到字节流
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"校勘记_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


async def _export_notes_markdown(
    notes: List[Dict],
    title: str,
    base_name: str,
    collation_names: List[str],
    include_statistics: bool,
    include_base_info: bool
) -> StreamingResponse:
    """导出Markdown格式校勘记"""
    lines = ["# 校勘记\n"]

    if include_base_info:
        lines.append("## 基本信息\n")
        lines.append(f"- **底本**：{base_name}")
        if collation_names:
            lines.append(f"- **校本**：{'、'.join(collation_names)}")
        lines.append(f"- **校勘记条数**：{len(notes)}")
        lines.append(f"- **生成时间**：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        lines.append("")

    if include_statistics:
        lines.append("## 统计分析\n")
        by_category = {}
        for note in notes:
            cat = note.get("category", "其他")
            by_category[cat] = by_category.get(cat, 0) + 1

        lines.append("### 异文类型分布\n")
        for category, count in by_category.items():
            lines.append(f"- {category}：{count}处")

        uncertain_count = sum(1 for n in notes if n.get("uncertain", False))
        if uncertain_count > 0:
            lines.append(f"\n**存疑条目**：{uncertain_count}处")
        lines.append("")

    lines.append("## 校勘记正文\n")
    for note in notes:
        formatted_text = note.get("formatted_text", "")
        if note.get("uncertain", False):
            lines.append(f"*{formatted_text}* （存疑）\n")
        else:
            lines.append(f"{formatted_text}\n")

    date_str = datetime.now().strftime('%Y年%m月%d日')
    footer_text = '\n---\n*本报告由"佛典标点与校勘研究平台"生成于' + date_str + '*'
    lines.append(footer_text)

    content = "\n".join(lines)
    buffer = BytesIO(content.encode('utf-8'))

    filename = f"校勘记_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    return StreamingResponse(
        buffer,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )


async def _export_notes_txt(
    notes: List[Dict],
    title: str,
    base_name: str,
    collation_names: List[str],
    include_statistics: bool,
    include_base_info: bool
) -> StreamingResponse:
    """导出TXT格式校勘记"""
    lines = ["校勘记", "=" * 40, ""]

    if include_base_info:
        lines.append("【基本信息】")
        lines.append(f"底本：{base_name}")
        if collation_names:
            lines.append(f"校本：{'、'.join(collation_names)}")
        lines.append(f"校勘记条数：{len(notes)}")
        lines.append(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        lines.append("")

    if include_statistics:
        lines.append("【统计分析】")
        by_category = {}
        for note in notes:
            cat = note.get("category", "其他")
            by_category[cat] = by_category.get(cat, 0) + 1

        for category, count in by_category.items():
            lines.append(f"  {category}：{count}处")

        uncertain_count = sum(1 for n in notes if n.get("uncertain", False))
        if uncertain_count > 0:
            lines.append(f"  存疑条目：{uncertain_count}处")
        lines.append("")

    lines.append("【校勘记正文】")
    lines.append("-" * 40)
    for note in notes:
        formatted_text = note.get("formatted_text", "")
        lines.append(formatted_text)

    lines.append("")
    lines.append("-" * 40)
    txt_date_str = datetime.now().strftime('%Y年%m月%d日')
    txt_footer = '本报告由"佛典标点与校勘研究平台"生成于' + txt_date_str
    lines.append(txt_footer)

    content = "\n".join(lines)
    buffer = BytesIO(content.encode('utf-8'))

    filename = f"校勘记_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    return StreamingResponse(
        buffer,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
    )
